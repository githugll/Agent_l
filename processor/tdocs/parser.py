"""Parse .docx/.doc/.docm files into plain text."""

import logging
import os
import re
from pathlib import Path

import docx

logger = logging.getLogger(__name__)


class TdocParser:
    """Extract text from 3GPP Tdoc documents (.docx, .doc, .docm)."""

    # Regex for section headers like "1.", "1.1", "2.3.1", "3.1.2.1"
    SECTION_PATTERN = re.compile(
        r"^(\d+(?:\.\d+){0,3})\s+\S"
    )

    def __init__(self, raw_dir: str, processed_dir: str):
        self.raw_dir = Path(raw_dir)
        self.texts_dir = Path(processed_dir) / "texts"
        self.texts_dir.mkdir(parents=True, exist_ok=True)
        self.failed_log = Path(processed_dir) / "failed_parsing.txt"

    def parse_docx(self, path: str) -> str:
        """Parse a .docx/.docm file and return plain text."""
        doc = docx.Document(path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n\n".join(paragraphs)

    def parse_doc(self, path: str) -> str:
        """Parse a legacy .doc file using textract."""
        try:
            import textract
            result = textract.process(path)
            return result.decode("utf-8", errors="replace")
        except Exception as e:
            logger.warning(f"textract failed for {path}: {e}, trying antiword")
            return self._parse_doc_antiword(path)

    def _parse_doc_antiword(self, path: str) -> str:
        """Fallback: parse .doc using antiword command line."""
        import subprocess
        try:
            result = subprocess.run(
                ["antiword", path],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0:
                return result.stdout
            logger.warning(f"antiword failed for {path}: {result.stderr}")
            return ""
        except FileNotFoundError:
            logger.warning("antiword not installed, skipping .doc file")
            return ""

    def _find_document_file(self, tdoc_dir: Path) -> Path | None:
        """Find the main document file in a Tdoc directory."""
        if not tdoc_dir.exists():
            return None

        files = list(tdoc_dir.iterdir())

        # Priority: .docx > .docm > .DOCX > .doc
        for ext in [".docx", ".docm"]:
            matches = [f for f in files if f.suffix.lower() == ext]
            if matches:
                # Prefer the file whose name starts with the Tdoc number
                tdoc_num = tdoc_dir.name
                primary = [f for f in matches if f.name.startswith(tdoc_num)]
                return primary[0] if primary else matches[0]

        # Check uppercase .DOCX
        docx_upper = [f for f in files if f.suffix == ".DOCX"]
        if docx_upper:
            return docx_upper[0]

        # Check .doc
        doc_files = [f for f in files if f.suffix.lower() == ".doc"]
        if doc_files:
            tdoc_num = tdoc_dir.name
            primary = [f for f in doc_files if f.name.startswith(tdoc_num)]
            return primary[0] if primary else doc_files[0]

        return None

    def extract_text(self, tdoc_number: str) -> dict | None:
        """Extract text for a single Tdoc. Returns dict with text and metadata."""
        # Find the Tdoc directory under raw/TSGR2_134/
        tdoc_dir = self.raw_dir / "TSGR2_134" / tdoc_number
        doc_file = self._find_document_file(tdoc_dir)

        if doc_file is None:
            logger.warning(f"No document file found for {tdoc_number}")
            return None

        try:
            ext = doc_file.suffix.lower()
            if ext in (".docx", ".docm"):
                text = self.parse_docx(str(doc_file))
            elif ext == ".doc":
                text = self.parse_doc(str(doc_file))
            elif ext == ".docx" and doc_file.suffix == ".DOCX":
                text = self.parse_docx(str(doc_file))
            else:
                logger.warning(f"Unsupported format: {doc_file}")
                return None

            if not text or len(text.strip()) < 50:
                logger.warning(f"Empty or very short text for {tdoc_number}")
                return None

            return {
                "tdoc_number": tdoc_number,
                "text": text,
                "source_file": str(doc_file),
                "char_count": len(text),
                "word_count": len(text.split()),
            }

        except Exception as e:
            logger.error(f"Failed to parse {tdoc_number}: {e}")
            return None

    def process_all(self, manifest_path: str) -> dict:
        """Process all Tdocs from manifest. Returns stats dict."""
        import json
        from tqdm import tqdm

        with open(manifest_path) as f:
            manifest = json.load(f)

        tdoc_numbers = [d["tdoc_number"] for d in manifest["documents"]]
        total = len(tdoc_numbers)

        success = 0
        failed = 0
        failed_list = []

        for tdoc_num in tqdm(tdoc_numbers, desc="Parsing Tdocs"):
            result = self.extract_text(tdoc_num)
            if result:
                # Save extracted text
                text_path = self.texts_dir / f"{tdoc_num}.txt"
                text_path.write_text(result["text"], encoding="utf-8")
                success += 1
            else:
                failed += 1
                failed_list.append(tdoc_num)

        # Write failed log
        if failed_list:
            with open(self.failed_log, "w") as f:
                for tdoc_num in failed_list:
                    f.write(f"{tdoc_num}\n")

        stats = {
            "total": total,
            "success": success,
            "failed": failed,
            "success_rate": success / total * 100,
            "failed_list": failed_list,
        }

        logger.info(
            f"Parsing complete: {success}/{total} "
            f"({stats['success_rate']:.1f}%), {failed} failed"
        )
        return stats
