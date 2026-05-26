"""Spec PDF parser: extracts structured text from 3GPP specification PDFs.

Uses pdfplumber as primary (better at tables/multi-column) with PyPDF2 fallback.
"""

import logging
import os
import re
from pathlib import Path

from tqdm import tqdm

logger = logging.getLogger(__name__)

# 3GPP spec page header/footer pattern to strip
HEADER_FOOTER_RE = re.compile(
    r"3GPP\s+TS\s+\d+\.\d+\s+[VRr]?\d+\.\d+.*?"
    r"|^\s*\d+\s*$"
    r"|^Vol\.\s+\d+.*?$"
    r"|^Chapter\s+\d+.*?$",
    re.MULTILINE | re.IGNORECASE,
)
WHITESPACE_RE = re.compile(r"[ \t]+")
MULTILINE_RE = re.compile(r"\n{3,}")


class SpecParser:
    """Extract text from 3GPP spec PDFs."""

    def __init__(self, raw_dir: str, processed_dir: str):
        self.raw_dir = Path(raw_dir)
        self.texts_dir = Path(processed_dir) / "texts"
        self.texts_dir.mkdir(parents=True, exist_ok=True)
        self.failed_log = Path(processed_dir) / "failed_parsing.txt"

    # ── Public API ────────────────────────────────────────────────────────────

    def parse_spec(self, spec_number: str, pdf_path: str) -> dict | None:
        """Parse a single spec document (PDF or DOCX). Returns dict with text, meta; saves .txt file."""
        spec_id = spec_number.replace(".", "")
        meta_path = self.texts_dir / f"{spec_id}_meta.json"

        # Skip if already parsed
        if meta_path.exists():
            txt_path = self.texts_dir / f"{spec_id}.txt"
            if txt_path.exists():
                with open(txt_path, encoding="utf-8") as f:
                    text = f.read()
                return {
                    "spec_number": spec_number,
                    "text": text,
                    "char_count": len(text),
                }

        # Extract based on file type
        ext = Path(pdf_path).suffix.lower()
        if ext == ".pdf":
            text = self._extract_with_pdfplumber(pdf_path)
            if not text or len(text) < 500:
                text = self._extract_with_pypdf2(pdf_path)
        elif ext in (".docx", ".docm"):
            text = self._extract_docx(pdf_path)
        elif ext == ".doc":
            text = self._extract_doc(pdf_path)
        else:
            logger.warning(f"Unsupported format: {pdf_path}")
            return None

        if not text or len(text) < 500:
            logger.error(f"Empty extraction for {spec_number}")
            return None

        # Clean
        text = self._clean_text(text, spec_number)
        logger.info(f"  Parsed {spec_number}: {len(text):,} chars")

        # Save
        txt_path = self.texts_dir / f"{spec_id}.txt"
        txt_path.write_text(text, encoding="utf-8")

        # Extract metadata
        meta = self._extract_metadata(text, spec_number)
        import json
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)

        return {
            "spec_number": spec_number,
            "text": text,
            "char_count": len(text),
            "meta": meta,
        }

    def process_all(self, manifest_path: str) -> dict:
        """Process all specs from manifest. Returns stats dict."""
        import json

        with open(manifest_path, encoding="utf-8") as f:
            manifest = json.load(f)

        specs = [s for s in manifest.get("specs", []) if s.get("status") == "extracted"]
        total = len(specs)
        success = 0
        failed = []
        total_chars = 0

        for entry in tqdm(specs, desc="Parsing specs"):
            result = self.parse_spec(entry["spec_number"], entry["pdf_path"])
            if result:
                success += 1
                total_chars += result["char_count"]
            else:
                failed.append(entry["spec_number"])

        if failed:
            with open(self.failed_log, "w") as f:
                for s in failed:
                    f.write(f"{s}\n")

        stats = {
            "total": total,
            "success": success,
            "failed": failed,
            "total_chars": total_chars,
        }
        logger.info(f"Parsing: {success}/{total} specs, {total_chars:,} total chars")
        return stats

    # ── PDF extraction ────────────────────────────────────────────────────────

    def _extract_with_pdfplumber(self, pdf_path: str) -> str:
        """Extract text with pdfplumber (preserves tables/columns)."""
        try:
            import pdfplumber
            texts = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    # Extract text (simple mode, no layout analysis needed)
                    t = page.extract_text()
                    if t:
                        texts.append(t)
                    # Also try table extraction for dense pages
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = self._table_to_text(table)
                            if table_text:
                                texts.append(table_text)
            return "\n\n".join(texts)
        except Exception as e:
            logger.warning(f"pdfplumber failed for {pdf_path}: {e}")
            return ""

    def _extract_with_pypdf2(self, pdf_path: str) -> str:
        """Fallback: PyPDF2 basic text extraction."""
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(pdf_path)
            texts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    texts.append(t)
            return "\n\n".join(texts)
        except Exception as e:
            logger.warning(f"PyPDF2 failed for {pdf_path}: {e}")
            return ""

    def _extract_docx(self, docx_path: str) -> str:
        """Extract text from .docx using python-docx."""
        try:
            import docx
            doc = docx.Document(docx_path)
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        paragraphs.append(" | ".join(row_texts))
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.warning(f"python-docx failed for {docx_path}: {e}")
            return ""

    def _extract_doc(self, doc_path: str) -> str:
        """Fallback: extract .doc using antiword."""
        import subprocess
        try:
            result = subprocess.run(
                ["antiword", doc_path],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0:
                return result.stdout
            logger.warning(f"antiword failed for {doc_path}: {result.stderr}")
            return ""
        except FileNotFoundError:
            logger.warning("antiword not installed, skipping .doc file")
            return ""

    def _table_to_text(self, table: list) -> str:
        """Convert a pdfplumber table to readable text rows."""
        lines = []
        for row in table:
            cells = [str(c).strip() if c else "" for c in row]
            line = " | ".join(cells)
            if line.strip():
                lines.append(line)
        return "\n".join(lines)

    # ── Text cleaning ─────────────────────────────────────────────────────────

    def _clean_text(self, text: str, spec_number: str) -> str:
        """Remove headers/footers, normalize whitespace."""
        # Remove header/footer lines
        lines = text.split("\n")
        cleaned = []
        for line in lines:
            # Skip lines that are just page artifacts
            if self._is_artificial_line(line):
                continue
            # Strip leading/trailing whitespace per line
            line = line.strip()
            if line:
                cleaned.append(line)
        text = "\n".join(cleaned)

        # Normalize whitespace
        text = WHITESPACE_RE.sub(" ", text)
        text = MULTILINE_RE.sub("\n\n", text)

        return text

    def _is_artificial_line(self, line: str) -> bool:
        """Detect page header/footer/artificial lines."""
        line = line.strip()
        if not line:
            return True
        # Common header/footer patterns
        if re.match(r"^\d+\s*$", line):          # Standalone page number
            return True
        if "3GPP TS" in line and len(line) < 120 and not line[0].isalnum():
            # Header line starting mid-sentence
            return True
        if re.match(r"^Vol\.\s+\d+", line, re.I):
            return True
        return False

    # ── Metadata extraction ───────────────────────────────────────────────────

    def _extract_metadata(self, text: str, spec_number: str) -> dict:
        """Extract title, version, release, date from cover page."""
        meta = {
            "spec_number": spec_number,
            "spec_id": spec_number.replace(".", ""),
        }

        # Try to extract title from first 2000 chars
        sample = text[:2000]
        title_match = re.search(
            r"(?:TS\s+)?(\d+\.\d+)\s*\n?\s*(.{5,120})",
            sample,
            re.MULTILINE,
        )
        if title_match:
            # Filter out the version line (e.g. "38.321 V19.2.0")
            title_text = title_match.group(2)
            if not re.match(r"^V\d+\.\d+", title_text, re.I):
                meta["title"] = title_text.strip()

        # Version
        version_match = re.search(r"V(\d+)\.(\d+)", sample)
        if version_match:
            meta["version"] = f"{version_match.group(1)}.{version_match.group(2)}"

        # Release
        if "Release 19" in sample or "Rel-19" in sample:
            meta["release"] = "Rel-19"
        elif "Release 17" in sample or "Rel-17" in sample:
            meta["release"] = "Rel-17"
        elif "Release 18" in sample or "Rel-18" in sample:
            meta["release"] = "Rel-18"

        # Date
        date_match = re.search(
            r"(\d{4}[-/]\d{2}[-/]\d{2})|(\d{2}\s+\w+\s+\d{4})",
            sample,
        )
        if date_match:
            meta["date"] = date_match.group(0)

        return meta
